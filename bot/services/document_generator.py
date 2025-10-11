from docx import Document
from docx.shared import Pt
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.enums import TA_LEFT


def create_docx(text: str, output_path: str):
    """Create a Word document with the summary text"""
    doc = Document()
    
    # Add title
    title = doc.add_heading('Summary', 0)
    
    # Add summary text
    paragraphs = text.split('\n')
    for para in paragraphs:
        if para.strip():
            p = doc.add_paragraph(para)
            # Set font
            for run in p.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(11)
    
    doc.save(output_path)


def create_pdf(text: str, output_path: str):
    """Create a PDF document with the summary text"""
    doc = SimpleDocTemplate(output_path, pagesize=A4,
                            rightMargin=72, leftMargin=72,
                            topMargin=72, bottomMargin=18)
    
    # Container for the 'Flowable' objects
    elements = []
    
    # Define styles
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name='Justify', alignment=TA_LEFT))
    
    # Add title
    title_style = styles['Heading1']
    elements.append(Paragraph("Summary", title_style))
    elements.append(Spacer(1, 12))
    
    # Add text paragraphs
    paragraphs = text.split('\n')
    for para in paragraphs:
        if para.strip():
            # Replace special characters that might cause issues
            para_text = para.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            p = Paragraph(para_text, styles['Normal'])
            elements.append(p)
            elements.append(Spacer(1, 12))
    
    # Build PDF
    doc.build(elements)

